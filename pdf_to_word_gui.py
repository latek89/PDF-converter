import os
import threading
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk
from pdf2docx import Converter
from spire.pdf.common import *
from spire.pdf import *
from docx import Document
import multiprocessing
import time

# Ustawienia estetyczne (Premium Look)
ctk.set_appearance_mode("dark")  # "light", "dark", "system"
ctk.set_default_color_theme("blue")

class PDFConverterApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("PDF → Word Professional Converter (GLOBAL CLEAN 1:1)")
        self.geometry("600x600")
        
        # Konfiguracja siatki
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(4, weight=1)

        # Nagłówek
        self.header_label = ctk.CTkLabel(self, text="Konwerter PDF do Word (1:1)", font=ctk.CTkFont(size=24, weight="bold"))
        self.header_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        self.subheader_label = ctk.CTkLabel(self, text="Usuwanie znaków wodnych z całego dokumentu", font=ctk.CTkFont(size=14))
        self.subheader_label.grid(row=1, column=0, padx=20, pady=(0, 20))

        # Tryb konwersji
        self.mode_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.mode_frame.grid(row=2, column=0, padx=20, pady=5, sticky="ew")
        self.mode_switch = ctk.CTkSwitch(self.mode_frame, text="Tryb Profesjonalny (Pełne Czyszczenie)", 
                                         font=ctk.CTkFont(size=13, weight="bold"))
        self.mode_switch.select()
        self.mode_switch.pack(pady=5)

        # Przyciski wyboru plików i folderu
        self.button_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.button_frame.grid(row=3, column=0, padx=20, pady=10, sticky="ew")
        self.button_frame.grid_columnconfigure((0, 1), weight=1)

        self.add_files_button = ctk.CTkButton(self.button_frame, text="Wybierz Pliki PDF", command=self.add_files, 
                                            height=40, font=ctk.CTkFont(size=14, weight="bold"))
        self.add_files_button.grid(row=0, column=0, padx=10, pady=5, sticky="ew")

        self.set_output_button = ctk.CTkButton(self.button_frame, text="Zmień folder zapisu", command=self.set_output_folder, 
                                            height=40, fg_color="gray", hover_color="#555555", font=ctk.CTkFont(size=14))
        self.set_output_button.grid(row=0, column=1, padx=10, pady=5, sticky="ew")

        self.convert_button = ctk.CTkButton(self, text="KONWERTUJ (WSZYSTKIE STRONY CZYSTE)", command=self.start_conversion, 
                                           fg_color="#2ecc71", hover_color="#27ae60", height=50, 
                                           font=ctk.CTkFont(size=18, weight="bold"))
        self.convert_button.grid(row=5, column=0, padx=30, pady=20, sticky="ew")
        self.convert_button.configure(state="disabled")

        # Lista plików
        self.output_label = ctk.CTkLabel(self, text="Folder zapisu: Domyślny", font=ctk.CTkFont(size=11), text_color="#aaaaaa")
        self.output_label.grid(row=4, column=0, padx=20, pady=(0, 10), sticky="n")

        self.file_list_label = ctk.CTkLabel(self, text="Wybrane pliki:", font=ctk.CTkFont(size=12, weight="bold"))
        self.file_list_label.grid(row=4, column=0, padx=20, pady=(30, 0), sticky="w")

        self.file_textbox = ctk.CTkTextbox(self, height=100)
        self.file_textbox.grid(row=4, column=0, padx=20, pady=(60, 10), sticky="nsew")
        self.file_textbox.configure(state="disabled")

        # Sekcja postępu
        self.progress_label = ctk.CTkLabel(self, text="Gotowy do pracy", font=ctk.CTkFont(size=12))
        self.progress_label.grid(row=6, column=0, padx=20, pady=(0, 5))

        self.progress_bar = ctk.CTkProgressBar(self, mode="determinate")
        self.progress_bar.grid(row=7, column=0, padx=20, pady=(0, 20), sticky="ew")
        self.progress_bar.set(0)

        # Zmienne logiczne
        self.selected_files = []
        self.output_folder = None
        self.is_converting = False

    def add_files(self):
        files = filedialog.askopenfilenames(
            title="Wybierz pliki PDF",
            filetypes=[("Pliki PDF", "*.pdf")]
        )
        if files:
            self.selected_files = list(files)
            self.update_file_list()
            self.convert_button.configure(state="normal")
            self.progress_label.configure(text=f"Wczytano {len(files)} plików.")

    def set_output_folder(self):
        folder = filedialog.askdirectory(title="Wybierz folder zapisu")
        if folder:
            self.output_folder = folder
            self.output_label.configure(text=f"Folder zapisu: {os.path.basename(folder)}")
            self.set_output_button.configure(fg_color="#3498db", hover_color="#2980b9")

    def update_file_list(self):
        self.file_textbox.configure(state="normal")
        self.file_textbox.delete("1.0", tk.END)
        for f in self.selected_files:
            self.file_textbox.insert(tk.END, f"{os.path.basename(f)}\n")
        self.file_textbox.configure(state="disabled")

    def remove_spire_watermark(self, docx_path):
        """Przeprowadza pełny skan dokumentu i usuwa znak wodny ze wszystkich stron, nagłówków i stopek."""
        try:
            doc = Document(docx_path)
            
            # SŁOWA KLUCZOWE DO USUNIĘCIA
            targets = ["Evaluation Warning", "Spire.PDF", "created with Spire.PDF"]

            def clean_blocks(blocks):
                for p in list(blocks):
                    if any(t in p.text for t in targets):
                        try:
                            # Usuń akapit z rodzica
                            p_element = p._element
                            p_element.getparent().remove(p_element)
                        except:
                            pass

            # 1. Przeskanuj wszystkie akapity w głównym dokumencie
            clean_blocks(doc.paragraphs)

            # 2. Przeskanuj nagłówki i stopki we wszystkich sekcjach
            for section in doc.sections:
                header = section.header
                if header:
                    clean_blocks(header.paragraphs)
                
                footer = section.footer
                if footer:
                    clean_blocks(footer.paragraphs)

            # 3. Przeskanuj komórki tabel (na wszelki wypadek)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        clean_blocks(cell.paragraphs)

            doc.save(docx_path)
        except Exception as e:
            print(f"Błąd globalnego usuwania znaku wodnego: {e}")

    def start_conversion(self):
        if not self.selected_files or self.is_converting:
            return

        self.is_converting = True
        self.toggle_ui("disabled")
        
        thread = threading.Thread(target=self.convert_process)
        thread.daemon = True
        thread.start()

    def toggle_ui(self, state):
        self.add_files_button.configure(state=state)
        self.set_output_button.configure(state=state)
        self.convert_button.configure(state=state)
        self.mode_switch.configure(state=state)

    def convert_process(self):
        total_files = len(self.selected_files)
        use_pro_mode = self.mode_switch.get()
        
        for i, pdf_path in enumerate(self.selected_files):
            try:
                base_name = os.path.basename(pdf_path)
                file_name_only = os.path.splitext(base_name)[0]
                
                if self.output_folder:
                    save_dir = self.output_folder
                else:
                    save_dir = os.path.dirname(pdf_path)
                
                docx_path = os.path.join(save_dir, f"{file_name_only}.docx")
                
                self.update_status(f"Przetwarzam ({i+1}/{total_files}): {base_name}...")
                
                if use_pro_mode:
                    # SILNIK PROFESJONALNY: Spire.Pdf
                    doc = PdfDocument()
                    doc.LoadFromFile(pdf_path)
                    doc.SaveToFile(docx_path, FileFormat.DOCX)
                    doc.Close()
                    
                    # CZYSZCZENIE GLOBALNE (Wszystkie strony i nagłówki)
                    self.remove_spire_watermark(docx_path)
                else:
                    cv = Converter(pdf_path)
                    cv.convert(docx_path, multi_processing=True, cpu_count=4)
                    cv.close()
                
                progress = (i + 1) / total_files
                self.update_progress(progress)

            except Exception as e:
                self.after(0, lambda err=e, p=pdf_path: messagebox.showerror("Błąd", f"Wystąpił problem: {os.path.basename(p)}\n{str(err)}"))

        self.update_status("Sukces! Dokumenty są czyste.")
        self.after(0, lambda: messagebox.showinfo("Sukces", "Konwersja 1:1 zakończona. Usunięto znaki wodne ze wszystkich stron!"))
        self.after(0, self.reset_ui)

    def update_status(self, text):
        self.after(0, lambda: self.progress_label.configure(text=text))

    def update_progress(self, val):
        self.after(0, lambda: self.progress_bar.set(val))

    def reset_ui(self):
        self.is_converting = False
        self.toggle_ui("normal")
        self.progress_bar.set(0)
        self.selected_files = []
        self.output_folder = None
        self.output_label.configure(text="Folder zapisu: Domyślny")
        self.set_output_button.configure(fg_color="gray", hover_color="#555555")
        self.update_file_list()
        self.convert_button.configure(state="disabled")

if __name__ == "__main__":
    multiprocessing.freeze_support()
    app = PDFConverterApp()
    app.mainloop()
